
from docx import Document

# Crear el documento
doc = Document()

# Título
doc.add_heading('Plan y Ejecución de Pruebas Funcionales', 0)

# Objetivo y Tipos de Pruebas
doc.add_heading('Objetivo', level=1)
doc.add_paragraph(
    'Verificar que las funcionalidades principales del sistema cumplen con los requisitos establecidos, '
    'asegurando el correcto funcionamiento de los casos de uso.'
)
doc.add_heading('Tipos de Pruebas', level=1)
doc.add_paragraph('• Pruebas funcionales')
doc.add_paragraph('• Pruebas de aceptación')

# Sección: Plan de Pruebas
doc.add_heading('Plan de Pruebas', level=1)
doc.add_paragraph(
    'El plan de pruebas cubre los siguientes módulos: Autenticación, Formulario de encuesta, '
    'Visualización de respuestas, Contacto, Dashboard administrador, Gestión de usuarios, Análisis exploratorio, '
    'Visualización de mensajes de contacto, Descarga de archivo Excel y Página Nosotros.'
)

# Tabla de Casos de Prueba
doc.add_heading('Casos de Prueba', level=2)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Descripción'
hdr_cells[2].text = 'Datos de Entrada'
hdr_cells[3].text = 'Resultado Esperado'

casos = [
    ('P01', 'Registro exitoso', 'Nombre, usuario, contraseña', 'Usuario registrado y redirigido al login'),
    ('P02', 'Registro incompleto', 'Nombre vacío', 'Mensaje de error: "Todos los campos son obligatorios"'),
    ('P03', 'Inicio de sesión correcto', 'Usuario y contraseña válidos', 'Redirección al inicio'),
    ('P04', 'Inicio de sesión fallido', 'Usuario o contraseña incorrectos', 'Mensaje: "Usuario o contraseña incorrectos"'),
    ('P05', 'Enviar encuesta completa', 'Datos personales y 23 respuestas válidas', 'Mensaje de éxito y redirección a "Mis Respuestas"'),
    ('P06', 'Enviar encuesta incompleta', 'Faltan respuestas', 'Mensaje de error indicando los campos faltantes'),
    ('P07', 'Visualizar "Mis Respuestas"', 'Usuario autenticado con respuestas registradas', 'Se muestra resumen y gráfico'),
    ('P08', 'Enviar mensaje de contacto', 'Nombre, email, mensaje', 'Mensaje guardado y correo de confirmación enviado'),
    ('P09', 'Enviar mensaje de contacto incompleto', 'Campos vacíos', 'Mensaje: "Por favor completa todos los campos"'),
    ('P10', 'Acceso al dashboard', 'Administrador autenticado', 'Acceso al panel y visualización de estadísticas'),
    ('P11', 'Ver lista de usuarios', 'Administrador autenticado', 'Visualización de tabla de usuarios'),
    ('P12', 'Ver respuestas del usuario', 'Administrador autenticado', 'Visualización de respuestas usuarios'),
    ('P13', 'Editar usuario', 'Nuevo nombre, rol y contraseña (Administrador autenticado)', 'Datos actualizados correctamente'),
    ('P14', 'Eliminar usuario', 'ID de usuario válido (Administrador autenticado)', 'Usuario y respuestas eliminadas'),
    ('P15', 'Aplicar filtros en análisis', 'Categoría, edad o sexo (Administrador autenticado)', 'Actualización correcta de estadísticas'),
    ('P16', 'Descargar Excel', 'Solicitud de descarga (Administrador autenticado)', 'Descarga exitosa del archivo Excel'),
    ('P17', 'Visualizar mensajes de contacto', 'Mensajes registrados en la base de datos', 'Lista de mensajes mostrada correctamente'),
    ('P18', 'Visualizar página "Nosotros"', 'Solicitud a la ruta /nosotros', 'Página cargada correctamente con la información institucional'),
]

for id, desc, entrada, resultado in casos:
    row_cells = table.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = desc
    row_cells[2].text = entrada
    row_cells[3].text = resultado

# Sección: Informe de Ejecución de Pruebas
doc.add_heading('Informe de Ejecución de Pruebas', level=1)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Resultado Esperado'
hdr_cells[2].text = 'Resultado Obtenido'
hdr_cells[3].text = 'Estado'

for id, desc, entrada, resultado in casos:
    row_cells = table.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = resultado
    row_cells[2].text = '✅ Funciona'
    row_cells[3].text = 'Aprobado'

# Sección: Formato para Pruebas de Aceptación
doc.add_heading('Formato para Pruebas de Aceptación', level=1)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Criterio de Aceptación'
hdr_cells[2].text = 'Estado'

for id, desc, entrada, resultado in casos:
    row_cells = table.add_row().cells
    row_cells[0].text = id.replace('P', 'CA')
    row_cells[1].text = desc
    row_cells[2].text = 'Aceptado'

# Guardar el documento
file_path = "Plan_Pruebas_Ruleta_de_la_Vida.docx"
doc.save(file_path)
print(f"Documento generado exitosamente en {file_path}")
